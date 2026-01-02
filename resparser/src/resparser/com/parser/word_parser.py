from ..parser.file_parser import FileParser
from docx import Document


class WordParser(FileParser):
    def parse(self, file_path: str) -> str:
        """
        Extract text from a Word (.docx) resume using python-docx.

        Returns:
            Full text content as a single string
        """
        document = Document(file_path)

        text_chunks = []

        # Paragraph text
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_chunks.append(text)

        # Tables (some resumes use tables for layout)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_chunks.append(cell_text)

        return "\n".join(text_chunks)

